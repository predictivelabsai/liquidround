"""Export endpoints — XLSX for tables, DOCX for memos, Plotly charts.

POST /app/export/xlsx   — table JSON → .xlsx download
POST /app/export/docx   — markdown  → .docx download
POST /app/chart          — table JSON → Plotly JSON spec
"""

from __future__ import annotations

import io
import json
import logging
import re
from datetime import datetime

from fasthtml.common import APIRouter
from starlette.requests import Request
from starlette.responses import Response, JSONResponse

ar = APIRouter()
log = logging.getLogger(__name__)

# LiquidRound palette
_LR_NAVY = "0B1220"
_LR_AMBER = "F59E0B"
_LR_SLATE = "64748B"
_LR_INK = "E5E7EB"


# ── XLSX export ────────────────────────────────────────────────────────

@ar("/app/export/xlsx", methods=["POST"])
async def export_xlsx(request: Request):
    form = await request.form()
    raw = form.get("data", "")
    try:
        data = json.loads(raw)
    except Exception:
        return JSONResponse({"error": "bad json"}, status_code=400)

    columns = data.get("columns", [])
    rows = data.get("rows", [])
    title = data.get("title", "LiquidRound Export")

    if not columns or not rows:
        return JSONResponse({"error": "empty data"}, status_code=400)

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]

    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color=_LR_NAVY, end_color=_LR_NAVY, fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="1E293B"),
        right=Side(style="thin", color="1E293B"),
        top=Side(style="thin", color="1E293B"),
        bottom=Side(style="thin", color="1E293B"),
    )

    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    for row_idx, row in enumerate(rows, 2):
        for col_idx, col_name in enumerate(columns, 1):
            val = row.get(col_name)
            if isinstance(val, str):
                try:
                    val = float(val.replace(",", ""))
                except (ValueError, AttributeError):
                    pass
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = thin_border

    for col_idx, col_name in enumerate(columns, 1):
        max_len = len(str(col_name))
        for row in rows[:20]:
            cell_val = str(row.get(col_name, ""))
            max_len = max(max_len, len(cell_val))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-")[:40] or "export"
    filename = f"liquidround-{safe_title}.xlsx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── DOCX export ───────────────────────────────────────────────────────

@ar("/app/export/docx", methods=["POST"])
async def export_docx(request: Request):
    form = await request.form()
    markdown = form.get("markdown", "")
    title = form.get("title", "LiquidRound Document")

    if not markdown:
        return JSONResponse({"error": "empty markdown"}, status_code=400)

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x0B, 0x12, 0x20)

    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x12, 0x20)

    doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}", style="Subtitle")

    in_table = False
    table_rows = []

    for line in markdown.split("\n"):
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            if re.fullmatch(r"\|[\s\-:|]+\|", stripped):
                continue
            cells = [c.strip() for c in stripped[1:-1].split("|")]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue

        if in_table:
            _add_table(doc, table_rows)
            in_table = False
            table_rows = []

        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("```"):
            continue
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    if in_table:
        _add_table(doc, table_rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-")[:40] or "document"
    filename = f"liquidround-{safe_title}.docx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_table(doc, rows):
    if not rows:
        return
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols, style="Table Grid")

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    p.style.font.size = Pt(9)
                if i == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn("w:shd"), {
                        qn("w:fill"): _LR_NAVY,
                        qn("w:val"): "clear",
                    })
                    shading.append(shading_elm)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_formatted_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ── Chart generation ──────────────────────────────────────────────────

@ar("/app/chart", methods=["POST"])
async def generate_chart(request: Request):
    form = await request.form()
    raw = form.get("data", "")
    try:
        data = json.loads(raw)
    except Exception:
        return JSONResponse({"error": "bad json"}, status_code=400)

    columns = data.get("columns", [])
    rows = data.get("rows", [])
    title = data.get("title", "")

    if not columns or not rows:
        return JSONResponse({"error": "empty data"}, status_code=400)

    import plotly.graph_objects as go

    chart_type, fig = _auto_chart(columns, rows, title)

    fig.update_layout(
        template="plotly_dark",
        font=dict(family="Inter, Calibri, sans-serif", size=12, color="#E5E7EB"),
        title=dict(font=dict(size=16, color="#F59E0B")),
        plot_bgcolor="#0B1220",
        paper_bgcolor="#111A2E",
        margin=dict(l=50, r=30, t=50, b=40),
        height=380,
    )

    return JSONResponse({
        "chart_type": chart_type,
        "plotly": json.loads(fig.to_json()),
    })


def _auto_chart(columns: list[str], rows: list[dict], title: str):
    import plotly.graph_objects as go

    date_col = _find_date_column(columns, rows)
    numeric_cols = _find_numeric_columns(columns, rows)

    if date_col and numeric_cols:
        x_vals = [r.get(date_col, "") for r in rows]
        if len(rows) >= 20:
            fig = go.Figure()
            for nc in numeric_cols[:4]:
                y_vals = [_to_num(r.get(nc)) for r in rows]
                fig.add_trace(go.Scatter(
                    x=x_vals, y=y_vals, name=nc, mode="lines",
                    fill="tozeroy" if len(numeric_cols) == 1 else None,
                    line=dict(width=2),
                ))
            fig.update_layout(title=title)
            return "area", fig
        else:
            fig = go.Figure()
            for nc in numeric_cols[:4]:
                y_vals = [_to_num(r.get(nc)) for r in rows]
                fig.add_trace(go.Bar(x=x_vals, y=y_vals, name=nc))
            fig.update_layout(title=title, barmode="group")
            return "bar", fig

    cat_col = _find_categorical_column(columns, rows)
    if cat_col and len(numeric_cols) == 1:
        labels = [str(r.get(cat_col, "")) for r in rows]
        values = [_to_num(r.get(numeric_cols[0])) for r in rows]
        if len(rows) <= 8:
            fig = go.Figure(data=[go.Pie(
                labels=labels, values=values, hole=0.35,
                textinfo="label+percent",
                marker=dict(colors=_lr_palette(len(rows))),
            )])
            fig.update_layout(title=title)
            return "pie", fig
        else:
            fig = go.Figure(data=[go.Treemap(
                labels=labels, values=values, parents=[""] * len(rows),
                marker=dict(colors=_lr_palette(len(rows))),
                textinfo="label+value+percent root",
            )])
            fig.update_layout(title=title)
            return "treemap", fig

    if cat_col and len(numeric_cols) >= 2:
        x_vals = [str(r.get(cat_col, "")) for r in rows]
        fig = go.Figure()
        for nc in numeric_cols[:4]:
            y_vals = [_to_num(r.get(nc)) for r in rows]
            fig.add_trace(go.Bar(x=x_vals, y=y_vals, name=nc))
        fig.update_layout(title=title, barmode="group")
        return "grouped_bar", fig

    if numeric_cols:
        x_vals = list(range(1, len(rows) + 1))
        fig = go.Figure()
        for nc in numeric_cols[:4]:
            y_vals = [_to_num(r.get(nc)) for r in rows]
            fig.add_trace(go.Bar(x=x_vals, y=y_vals, name=nc))
        fig.update_layout(title=title, barmode="group")
        return "bar", fig

    return "bar", go.Figure()


def _find_date_column(columns, rows):
    date_patterns = re.compile(
        r"^\d{4}[-/]\d{2}([-/]\d{2})?$|^\d{2}[-/]\d{2}[-/]\d{4}$|"
        r"^(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)", re.I
    )
    date_keywords = {"date", "month", "year", "period", "quarter", "day", "time", "week"}
    for col in columns:
        if col.lower() in date_keywords:
            return col
        sample = [str(rows[i].get(col, "")) for i in range(min(3, len(rows)))]
        if all(date_patterns.match(s) for s in sample if s):
            return col
    return None


def _find_numeric_columns(columns, rows):
    numeric = []
    for col in columns:
        vals = [rows[i].get(col) for i in range(min(5, len(rows)))]
        num_count = sum(1 for v in vals if _is_numeric(v))
        if num_count >= len(vals) * 0.6:
            numeric.append(col)
    return numeric


def _find_categorical_column(columns, rows):
    for col in columns:
        vals = [rows[i].get(col) for i in range(min(5, len(rows)))]
        if all(isinstance(v, str) and not _is_numeric(v) for v in vals if v is not None):
            return col
    return None


def _is_numeric(v):
    if isinstance(v, (int, float)):
        return True
    if isinstance(v, str):
        try:
            float(v.replace(",", ""))
            return True
        except (ValueError, AttributeError):
            return False
    return False


def _to_num(v):
    if isinstance(v, (int, float)):
        return v
    if isinstance(v, str):
        try:
            return float(v.replace(",", ""))
        except (ValueError, AttributeError):
            return 0
    return 0


def _lr_palette(n):
    base = [
        "#F59E0B", "#3B82F6", "#10B981", "#EF4444",
        "#8B5CF6", "#EC4899", "#14B8A6", "#F97316",
        "#6366F1", "#84CC16",
    ]
    return (base * ((n // len(base)) + 1))[:n]
